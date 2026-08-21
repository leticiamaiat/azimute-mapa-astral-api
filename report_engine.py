# -*- coding: utf-8 -*-
"""
Motor de geracao do Mapa Astral Personalizado (.docx), no sistema visual
extraido do PDF-modelo (Barbara): Helvetica/Arial, navy #0a213b, dourado
#b79242, cabecalho/rodape centralizados, caixas de destaque com borda fina,
tabelas com cabecalho navy solido, roda natal tecnica em SVG/PNG.

Uso:
    from report_engine import build_docx_bytes
    docx_bytes = build_docx_bytes(data)   # data segue o esquema em SCHEMA.md / exemplo_carla.json
"""
import io
import math
import os
import re
import tempfile
import textwrap

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe

NAVY = RGBColor(0x0a, 0x21, 0x3b)
GOLD = RGBColor(0xb7, 0x92, 0x42)
GOLD_LINE = "d9c68a"
BODY_INK = RGBColor(0x30, 0x36, 0x3b)
SUB_INK = RGBColor(0x5a, 0x61, 0x68)
FOOTER_INK = RGBColor(0x77, 0x77, 0x77)
WHITE = RGBColor(0xff, 0xff, 0xff)
BLACK = RGBColor(0x11, 0x11, 0x11)
FONT = "Arial"

SIGNS = ["Áries", "Touro", "Gêmeos", "Câncer", "Leão", "Virgem", "Libra",
         "Escorpião", "Sagitário", "Capricórnio", "Aquário", "Peixes"]

PLANET_COLORS = {
    "Sol": "#e8a33d", "Lua": "#9fd3ef", "Mercúrio": "#7fd1ae", "Vênus": "#f2a6c9",
    "Marte": "#e15b5b", "Júpiter": "#c9a4f2", "Saturno": "#9aa0a8", "Urano": "#6fd1d1",
    "Netuno": "#6f8ee0", "Plutão": "#b06f3f", "Nodo Norte": "#d9c65a",
}

ZODIAC_GLYPH = {
    "Áries": "♈", "Touro": "♉", "Gêmeos": "♊", "Câncer": "♋", "Leão": "♌", "Virgem": "♍",
    "Libra": "♎", "Escorpião": "♏", "Sagitário": "♐", "Capricórnio": "♑", "Aquário": "♒", "Peixes": "♓",
}
PLANET_GLYPH = {
    "Sol": "☉", "Lua": "☽", "Mercúrio": "☿", "Vênus": "♀", "Marte": "♂", "Júpiter": "♃",
    "Saturno": "♄", "Urano": "♅", "Netuno": "♆", "Plutão": "♇", "Nodo Norte": "☊",
}

CARD_BG = "#faf6ea"
CARD_INK = "#20242a"
CARD_SUB = "#5c6470"
COVER_NAVY = "#0a1a30"
COVER_GOLD = "#b79242"
COVER_GOLD_LIGHT = "#e9c98a"

# ---------------------------------------------------------------- xml helpers

def set_cell_border(cell, color="b79242", sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)

def set_row_bottom_border(cell, color="e6e6e6", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    el = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
    borders.append(el)
    tcPr.append(borders)

def para_bottom_border(paragraph, color="d9c68a", sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_page_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.name = FONT; run.font.size = Pt(6.3); run.font.color.rgb = FOOTER_INK

# --------------------------------------------------------------- text helpers

BOLD_RE = re.compile(r'\*\*(.+?)\*\*')

def add_richtext(paragraph, text, size=9.2, color=BODY_INK, base_bold=False, italic=False, font=FONT):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color; r.bold = base_bold; r.italic = italic
        r = paragraph.add_run(m.group(1))
        r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color; r.bold = True; r.italic = italic
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color; r.bold = base_bold; r.italic = italic

def add_para(doc, text, size=9.2, color=BODY_INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=9, space_before=0, font=FONT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.32
    add_richtext(p, text, size=size, color=color, base_bold=bold, italic=italic, font=font)
    return p

# ----------------------------------------------------------- cover art helpers

def _track(text, sep=" "):
    """Insere espacamento entre letras para simular tracking tipografico em
    titulos curtos em caixa alta."""
    return sep.join(text)

def _dm(value, pos, neg):
    """Converte grau decimal (lat/lon) em graus-minutos com hemisferio, ex:
    -15.583 -> 15°35'S."""
    hemi = pos if value >= 0 else neg
    value = abs(value)
    d = int(value)
    m = int(round((value - d) * 60))
    if m == 60:
        m = 0
        d += 1
    return f"{d:02d}°{m:02d}'{hemi}"

def _fmt_date_short(date_str):
    try:
        y, m, d = date_str.split("-")
        return f"{int(d):02d}/{int(m):02d}/{y}"
    except Exception:
        return date_str or ""

def _fit_fontsize(fig, text_obj, max_width_frac, min_size=11):
    """Reduz o fontsize de `text_obj` ate caber em `max_width_frac` da largura
    da figura, sem quebrar linha (usado para nomes longos no titulo da capa)."""
    renderer = fig.canvas.get_renderer()
    fig_w_px = fig.bbox.width
    for _ in range(14):
        bbox = text_obj.get_window_extent(renderer=renderer)
        width_frac = bbox.width / fig_w_px
        if width_frac <= max_width_frac:
            break
        size = text_obj.get_fontsize()
        new_size = max(min_size, size * (max_width_frac / width_frac) * 0.98)
        if new_size >= size:
            break
        text_obj.set_fontsize(new_size)
        fig.canvas.draw()
    return text_obj.get_fontsize()

def _wrap_lines(text, width, max_lines=None):
    lines = textwrap.wrap((text or "").strip(), width=width) or [""]
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1].rstrip() + "…"
    return lines

# ------------------------------------------------------------- block renderers

def render_head(doc, eyebrow, title, subtitle):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(eyebrow); r.font.name = FONT; r.font.size = Pt(7.2); r.font.bold = True; r.font.color.rgb = GOLD

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title); r.font.name = FONT; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle); r.font.name = FONT; r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = SUB_INK

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("✦"); r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = GOLD

def render_callout(doc, heading, text):
    table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = True
    cell = table.cell(0, 0)
    set_cell_border(cell, color="b79242", sz=8)
    set_cell_margins(cell, top=90, bottom=90, left=160, right=160)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    r = cell.paragraphs[0].add_run(heading.upper())
    r.font.name = FONT; r.font.size = Pt(8.3); r.font.bold = True; r.font.color.rgb = NAVY
    p2 = cell.add_paragraph(); p2.paragraph_format.line_spacing = 1.3
    add_richtext(p2, text, size=7.8, color=BODY_INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def render_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        row = table.add_row(); c0, c1 = row.cells
        set_row_bottom_border(c0); set_row_bottom_border(c1)
        set_cell_margins(c0, top=60, bottom=60, left=0, right=100); set_cell_margins(c1, top=60, bottom=60, left=0, right=0)
        c0.width = Cm(6.5)
        r0 = c0.paragraphs[0].add_run(k); r0.font.name = FONT; r0.font.size = Pt(7.6); r0.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r1 = c1.paragraphs[0].add_run(v); r1.font.name = FONT; r1.font.size = Pt(7.6); r1.font.bold = True; r1.font.color.rgb = BLACK

def render_data_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; shade_cell(cell, "0a213b"); set_cell_margins(cell, top=70, bottom=70, left=110, right=110)
        r = cell.paragraphs[0].add_run(h); r.font.name = FONT; r.font.size = Pt(7.4); r.font.bold = True; r.font.color.rgb = WHITE
    for row_vals in rows:
        row = table.add_row()
        for i, val in enumerate(row_vals):
            cell = row.cells[i]; set_row_bottom_border(cell); set_cell_margins(cell, top=55, bottom=55, left=110, right=110)
            r = cell.paragraphs[0].add_run(str(val)); r.font.name = FONT; r.font.size = Pt(7.4); r.font.color.rgb = BLACK

def render_list(doc, items):
    for lead, rest in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.25
        r = p.add_run("•  "); r.font.name = FONT; r.font.size = Pt(8.6); r.font.color.rgb = GOLD
        r = p.add_run(lead + " "); r.font.name = FONT; r.font.size = Pt(8.6); r.font.bold = True; r.font.color.rgb = NAVY
        r = p.add_run(rest); r.font.name = FONT; r.font.size = Pt(8.6); r.font.color.rgb = BODY_INK

def render_subp(doc, items):
    """Bloco de subsecoes curtas: cada item e (titulo, texto). Renderiza uma
    linha-titulo em negrito (ex.: 'Sol em Escorpiao - essencia e verdade
    interior') seguida do paragrafo correspondente, no estilo dos relatorios
    de referencia (varias subsecoes curtas por pagina em vez de 1 paragrafo denso)."""
    for heading, text in items:
        if not (text and text.strip()):
            continue
        if heading and heading.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(heading.strip())
            r.font.name = FONT; r.font.size = Pt(9.2); r.font.bold = True; r.font.color.rgb = NAVY
        add_para(doc, text, space_before=0, space_after=10)

def render_bullets(doc, heading, items):
    """Bloco 'Sintese pratica': titulo curto + 3 marcadores diretos, sem caixa
    com borda (ao contrario de render_callout), no estilo dos relatorios de
    referencia usados como modelo."""
    items = [it for it in items if it and it.strip()]
    if not items:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(heading.upper())
    r.font.name = FONT; r.font.size = Pt(8.3); r.font.bold = True; r.font.color.rgb = NAVY
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run("•  "); r.font.name = FONT; r.font.size = Pt(8.6); r.font.color.rgb = GOLD
        r = p.add_run(text.strip()); r.font.name = FONT; r.font.size = Pt(8.6); r.font.color.rgb = BODY_INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def render_numlist(doc, items):
    for n, heading, text in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{n}. {heading}"); r.font.name = FONT; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = GOLD
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(10); p2.paragraph_format.line_spacing = 1.3
        add_richtext(p2, text, size=9.0, color=BODY_INK)

def render_quote(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = NAVY

def render_image(doc, path, width_cm):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))

def render_blocks(doc, blocks):
    for b in blocks:
        kind = b[0]
        if kind == 'p':
            if b[1] and b[1].strip():
                add_para(doc, b[1])
        elif kind == 'callout':
            if b[2] and b[2].strip():
                render_callout(doc, b[1], b[2])
        elif kind == 'kv': render_kv_table(doc, b[1])
        elif kind == 'table': render_data_table(doc, b[1], b[2])
        elif kind == 'list': render_list(doc, b[1])
        elif kind == 'subp': render_subp(doc, b[1])
        elif kind == 'bullets': render_bullets(doc, b[1], b[2])
        elif kind == 'numlist': render_numlist(doc, b[1])
        elif kind == 'quote': render_quote(doc, b[1])
        elif kind == 'image': render_image(doc, b[1], b[2])
        elif kind == 'linelist':
            for line in b[1]:
                add_para(doc, line, size=8.6, space_after=4)

# --------------------------------------------------------------- natal wheel

def lon(sign, deg, minute):
    return SIGNS.index(sign) * 30 + deg + minute / 60.0

def _draw_wheel(ax, planets, houses, ascendant, bg="#0a213b", dot_r=0.028,
                 show_signs=False, show_house_nums=False, use_glyphs=False, aspects=None):
    """Desenha a roda natal dentro de `ax` (axes ja configurado, xlim/ylim
    simetricos o bastante para caber raio ~1.0-1.3)."""
    asc = lon(ascendant["sign"], ascendant["deg"], ascendant["min"])

    def xy(r, longitude):
        rel = (longitude - asc) % 360
        theta = math.radians(180 + rel)
        return r * math.cos(theta), r * math.sin(theta)

    if bg:
        ax.add_patch(mpatches.Circle((0, 0), 1.0, facecolor=bg, edgecolor="none", zorder=0))

    R_OUT, R_IN = 1.0, 0.32
    ax.add_patch(mpatches.Circle((0, 0), R_OUT, fill=False, edgecolor="#b79242", linewidth=1.4))
    ax.add_patch(mpatches.Circle((0, 0), R_IN, fill=False, edgecolor="#b79242", linewidth=1.4))
    ax.add_patch(mpatches.Circle((0, 0), (R_OUT + R_IN) / 2, fill=False, edgecolor="#b79242", linewidth=0.6, alpha=0.5))

    for h in houses:
        L = lon(h["sign"], h["deg"], h["min"])
        x1, y1 = xy(R_IN, L); x2, y2 = xy(R_OUT, L)
        ax.plot([x1, x2], [y1, y2], color="#b79242", linewidth=0.7, alpha=0.8)

    if show_signs:
        for i, sign in enumerate(SIGNS):
            x, y = xy(R_OUT + 0.085, i * 30 + 15)
            ax.text(x, y, ZODIAC_GLYPH[sign], ha="center", va="center", fontsize=12.5,
                    color="#e9c98a", family="DejaVu Sans", zorder=4)

    if show_house_nums:
        cusp_lons = [lon(h["sign"], h["deg"], h["min"]) for h in houses]
        for i in range(12):
            start = cusp_lons[i]; end = cusp_lons[(i + 1) % 12]
            mid = start + ((end - start) % 360) / 2
            x, y = xy(R_IN * 0.62, mid)
            n = houses[i].get("n", i + 1)
            ax.text(x, y, str(n), ha="center", va="center", fontsize=7,
                    color="#8ea0bd", family=FONT, zorder=4)

    if aspects:
        def ppos(pname, r):
            p = next((x for x in planets if x["name"] == pname), None)
            if not p:
                return None
            return xy(r, lon(p["sign"], p["deg"], p["min"]))
        r_asp = R_IN * 0.98
        for a in aspects:
            pa, pb = ppos(a.get("a"), r_asp), ppos(a.get("b"), r_asp)
            if not pa or not pb:
                continue
            color = "#5f8fe0" if a.get("harmonico") else "#d6584e"
            ax.plot([pa[0], pb[0]], [pa[1], pb[1]], color=color, linewidth=1.1, alpha=0.8, zorder=2)

    radii_cycle = [R_IN + (R_OUT - R_IN) * 0.30, R_IN + (R_OUT - R_IN) * 0.55, R_IN + (R_OUT - R_IN) * 0.80]
    used = {}
    for pl in planets:
        L = lon(pl["sign"], pl["deg"], pl["min"])
        bucket = round(L / 18)
        slot = used.get(bucket, 0); used[bucket] = slot + 1
        rr = radii_cycle[slot % len(radii_cycle)]
        x, y = xy(rr, L)
        color = PLANET_COLORS.get(pl["name"], "#cccccc")
        if use_glyphs and pl["name"] in PLANET_GLYPH:
            ax.text(x, y, PLANET_GLYPH[pl["name"]], ha="center", va="center", fontsize=dot_r * 340,
                    color=color, family="DejaVu Sans", zorder=5,
                    path_effects=[pe.withStroke(linewidth=1.6, foreground="#0a1a30")])
        else:
            ax.add_patch(mpatches.Circle((x, y), dot_r, facecolor=color, edgecolor="#0a213b", linewidth=0.8, zorder=5))

def build_wheel_png(planets, houses, ascendant, out_path):
    """planets: list of dict(name, sign, deg, min); houses: list of dict(n, sign, deg, min);
    ascendant: dict(sign, deg, min)."""
    fig, ax = plt.subplots(figsize=(5.2, 5.2), dpi=300)
    fig.patch.set_facecolor("#0a213b"); ax.set_facecolor("#0a213b")
    ax.set_xlim(-1.3, 1.3); ax.set_ylim(-1.3, 1.3); ax.set_aspect("equal"); ax.axis("off")
    _draw_wheel(ax, planets, houses, ascendant, bg=None)
    fig.savefig(out_path, facecolor=fig.get_facecolor(), bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)
    return out_path

_MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
          "agosto", "setembro", "outubro", "novembro", "dezembro"]

def _format_birth_line(birth):
    try:
        y, m, d = birth["date"].split("-")
        date_fmt = f"{int(d):02d} de {_MESES[int(m)-1]} de {y}"
    except Exception:
        date_fmt = birth.get("date", "")
    time_fmt = birth.get("time", "")
    return f"{date_fmt} • {time_fmt}h • {birth.get('place','')}"

def _photo_rgba_cover(photo_path, box_w_in, box_h_in, dpi, fade_right=0.0, fade_bottom=0.0, fade_left=0.0,
                       warm_strength=0.16):
    """Abre a foto, corta no formato 'cover' (preenche a caixa sem distorcer),
    aplica um leve tom dourado quente (para casar com a paleta navy/dourada da
    capa) e um degrade de transparencia nas bordas indicadas para a imagem se
    fundir organicamente no fundo navy, em vez de parecer colada."""
    from PIL import Image, ImageOps
    import numpy as np
    img = Image.open(photo_path)
    img = ImageOps.exif_transpose(img).convert("RGBA")

    target_w, target_h = int(box_w_in * dpi), int(box_h_in * dpi)
    src_w, src_h = img.size
    target_ratio, src_ratio = target_w / target_h, src_w / src_h
    if src_ratio > target_ratio:
        new_h = src_h
        new_w = int(src_ratio and target_ratio * src_h)
        x0 = (src_w - new_w) // 2
        img = img.crop((x0, 0, x0 + new_w, src_h))
    else:
        new_w = src_w
        new_h = int(src_w / target_ratio)
        y0 = max(0, (src_h - new_h) // 3)
        img = img.crop((0, y0, src_w, min(src_h, y0 + new_h)))
    img = img.resize((target_w, target_h), Image.LANCZOS)

    arr = np.array(img).astype(np.float64)
    if warm_strength > 0:
        warm = np.array([214.0, 168.0, 96.0])
        arr[:, :, :3] = arr[:, :, :3] * (1 - warm_strength) + warm * warm_strength
    alpha = arr[:, :, 3] / 255.0
    if fade_right > 0:
        w = arr.shape[1]
        fw = int(w * fade_right)
        if fw > 0:
            grad = np.linspace(1, 0, fw)
            alpha[:, w - fw:] *= grad[np.newaxis, :]
    if fade_left > 0:
        w = arr.shape[1]
        fw = int(w * fade_left)
        if fw > 0:
            grad = np.linspace(0, 1, fw)
            alpha[:, :fw] *= grad[np.newaxis, :]
    if fade_bottom > 0:
        h = arr.shape[0]
        fh = int(h * fade_bottom)
        if fh > 0:
            grad = np.linspace(1, 0, fh)
            alpha[h - fh:, :] *= grad[:, np.newaxis]
    arr[:, :, 3] = alpha * 255.0
    return arr / 255.0

def build_placeholder_cover_png(name, birth, planets, houses, ascendant, out_path,
                                 photo_path=None, sections=None, aspects=None):
    """Capa navy/dourada no layout do modelo de referencia (Barbara): titulo,
    nome, triade em destaque, foto + roda natal tecnica lado a lado sobre o
    fundo estrelado, e um cartao claro com frase de abertura, 3 blocos de
    leitura rapida, tabela de posicoes planetarias e dados tecnicos de
    nascimento. `sections` (opcional): dict de textos gerados por IA com as
    chaves capa_abertura/capa_personalidade/capa_relacionamentos/
    capa_crescimento/capa_fechamento; na ausencia, usa textos de reserva a
    partir da triade. `aspects`: lista de chart_engine.compute_aspects, usada
    para desenhar as linhas de aspecto dentro da roda."""
    sections = sections or {}
    fig_w, fig_h = 8.27, 11.69  # A4 portrait
    dpi = 210
    fig = plt.figure(figsize=(fig_w, fig_h), dpi=dpi)
    fig.patch.set_facecolor(COVER_NAVY)

    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.set_facecolor(COVER_NAVY)

    def pfind(nm):
        return next((p for p in planets if p["name"] == nm), None)
    sol, lua = pfind("Sol"), pfind("Lua")
    mc = pfind("Meio do Céu")
    first_name = (name or "").split()[0] if name else "você"

    rng = __import__("random").Random(hash(name) % (2**31))
    for _ in range(140):
        sx, sy = rng.uniform(0.02, 0.98), rng.uniform(0.02, 0.98)
        ax.plot(sx, sy, marker="*", markersize=rng.uniform(1.0, 3.0),
                color="#c9a24b", alpha=rng.uniform(0.12, 0.5), zorder=1)

    # ---- moldura ornamental ao redor de toda a capa (dupla linha dourada + florao nos cantos)
    FRAME_M = 0.018
    ax.add_patch(mpatches.Rectangle((FRAME_M, FRAME_M), 1 - 2 * FRAME_M, 1 - 2 * FRAME_M,
                                     fill=False, edgecolor=COVER_GOLD, linewidth=1.3, alpha=0.9, zorder=3))
    ax.add_patch(mpatches.Rectangle((FRAME_M + 0.008, FRAME_M + 0.008 * fig_w / fig_h),
                                     1 - 2 * (FRAME_M + 0.008), 1 - 2 * (FRAME_M + 0.008 * fig_w / fig_h),
                                     fill=False, edgecolor=COVER_GOLD, linewidth=0.5, alpha=0.6, zorder=3))
    for cx, cy in ((FRAME_M, FRAME_M), (1 - FRAME_M, FRAME_M), (FRAME_M, 1 - FRAME_M), (1 - FRAME_M, 1 - FRAME_M)):
        ax.text(cx, cy, "✦", ha="center", va="center", fontsize=9, color=COVER_GOLD_LIGHT,
                family="DejaVu Sans", zorder=4)

    # ---- decoracoes de canto: estrela grande a esquerda, lua crescente a direita
    ax.text(0.045, 0.978, "✦", ha="center", va="center", fontsize=15, color=COVER_GOLD_LIGHT, zorder=3)
    ax.text(0.075, 0.965, "✧", ha="center", va="center", fontsize=9, color=COVER_GOLD, zorder=3)
    ax.add_patch(mpatches.Circle((0.947, 0.975), 0.016, facecolor=COVER_GOLD_LIGHT, edgecolor="none", zorder=3))
    ax.add_patch(mpatches.Circle((0.953, 0.978), 0.016, facecolor=COVER_NAVY, edgecolor="none", zorder=4))

    # ------------------------------------------------------------ cabecalho
    ax.text(0.5, 0.955, _track("MAPA ASTRAL DE NASCIMENTO"), ha="center", va="center",
            fontsize=14.5, color=COVER_GOLD_LIGHT, fontweight="bold", family="serif", zorder=3)
    ax.plot([0.40, 0.475], [0.936, 0.936], color=COVER_GOLD, linewidth=0.7, alpha=0.8, zorder=3)
    ax.text(0.5, 0.936, "✦", ha="center", va="center", fontsize=8, color=COVER_GOLD, zorder=3)
    ax.plot([0.525, 0.60], [0.936, 0.936], color=COVER_GOLD, linewidth=0.7, alpha=0.8, zorder=3)

    name_txt = ax.text(0.5, 0.906, name.upper(), ha="center", va="center", fontsize=23,
                        color=COVER_GOLD_LIGHT, fontweight="bold", family="serif", zorder=3)
    fig.canvas.draw()
    _fit_fontsize(fig, name_txt, max_width_frac=0.88, min_size=13)

    ax.text(0.5, 0.878, _format_birth_line(birth).upper(), ha="center", va="center",
            fontsize=10.5, color="#e7ecf3", family="serif", zorder=3)

    # ------------------------------------------------------------- triade
    ax.plot([0.30, 0.40], [0.833, 0.833], color=COVER_GOLD, linewidth=0.7, alpha=0.8, zorder=3)
    ax.text(0.5, 0.833, _track("TRÍADE PRINCIPAL"), ha="center", va="center", fontsize=10.5,
            color=COVER_GOLD, fontweight="bold", family="serif", zorder=3)
    ax.plot([0.60, 0.70], [0.833, 0.833], color=COVER_GOLD, linewidth=0.7, alpha=0.8, zorder=3)

    triade_cols = [
        ("☉", "SOL", sol["sign"] if sol else "-"),
        ("☽", "LUA", lua["sign"] if lua else "-"),
        (ZODIAC_GLYPH.get(ascendant["sign"], "✦"), "ASCENDENTE", ascendant["sign"]),
    ]
    for cx, (glyph, label, sign) in zip((0.27, 0.5, 0.73), triade_cols):
        ax.text(cx, 0.807, label, ha="center", va="center", fontsize=8.5,
                color=COVER_GOLD, family=FONT, fontweight="bold", zorder=3)
        ax.text(cx, 0.778, glyph, ha="center", va="center", fontsize=22,
                color="#ffffff", family="DejaVu Sans", zorder=3)
        ax.text(cx, 0.750, sign.upper(), ha="center", va="center", fontsize=10.5,
                color=COVER_GOLD_LIGHT, family="serif", fontweight="bold", zorder=3)

    # -------------------------------------------------- foto + roda natal
    HERO_TOP, HERO_BOTTOM = 0.718, 0.415
    if photo_path:
        photo_box = (0.045, HERO_BOTTOM, 0.435, HERO_TOP - HERO_BOTTOM)
        px0, py0, pw, ph = photo_box
        try:
            ax.add_patch(mpatches.Circle((px0 + pw * 0.5, py0 + ph * 0.92), 0.30,
                                          facecolor=COVER_GOLD, alpha=0.10, edgecolor="none", zorder=1))
            rgba = _photo_rgba_cover(photo_path, pw * fig_w, ph * fig_h, dpi, fade_right=0.22, fade_bottom=0.12)
            ax.imshow(rgba, extent=[px0, px0 + pw, py0, py0 + ph], zorder=2, aspect="auto")
        except Exception:
            photo_path = None
    if not photo_path:
        deco_ax = fig.add_axes([0.045, HERO_BOTTOM, 0.435, (HERO_TOP - HERO_BOTTOM)])
        deco_ax.set_xlim(-1, 1); deco_ax.set_ylim(-1, 1); deco_ax.axis("off"); deco_ax.patch.set_alpha(0)
        for _ in range(40):
            sx, sy = rng.uniform(-0.95, 0.95), rng.uniform(-0.95, 0.95)
            deco_ax.plot(sx, sy, marker="*", markersize=rng.uniform(1.5, 4.5),
                         color=COVER_GOLD_LIGHT, alpha=rng.uniform(0.2, 0.6))
        deco_ax.text(0, 0, ZODIAC_GLYPH.get(sol["sign"], "✦") if sol else "✦", ha="center", va="center",
                     fontsize=90, color=COVER_GOLD, alpha=0.5, family="DejaVu Sans")

    wheel_w = 0.45
    wheel_ax = fig.add_axes([0.505, HERO_BOTTOM, wheel_w, wheel_w * fig_w / fig_h])
    wheel_ax.set_xlim(-1.28, 1.28); wheel_ax.set_ylim(-1.28, 1.28)
    wheel_ax.set_aspect("equal"); wheel_ax.axis("off"); wheel_ax.patch.set_alpha(0)
    _draw_wheel(wheel_ax, planets, houses, ascendant, bg=None, dot_r=0.040,
                show_signs=True, show_house_nums=True, use_glyphs=True, aspects=aspects)

    # -------------------------------------------------------- cartao claro
    CARD_TOP, CARD_BOTTOM = 0.405, 0.148
    CARD_X0, CARD_X1 = 0.035, 0.965
    card = mpatches.FancyBboxPatch((CARD_X0, CARD_BOTTOM), CARD_X1 - CARD_X0, CARD_TOP - CARD_BOTTOM,
                                    boxstyle="round,pad=0,rounding_size=0.012",
                                    facecolor=CARD_BG, edgecolor=COVER_GOLD, linewidth=1.1, zorder=2)
    ax.add_patch(card)

    pad = 0.018
    inner_x0, inner_x1 = CARD_X0 + pad, CARD_X1 - pad
    cursor_y = CARD_TOP - pad

    # frase de abertura, em caixa com borda fina
    abertura = sections.get("capa_abertura") or (
        f'Você nasceu para transformar {sol["sign"].lower() if sol else "sua essência"} em direção, '
        f'e {lua["sign"].lower() if lua else "sua sensibilidade"} em presença.')
    quote_h = 0.052
    ax.add_patch(mpatches.FancyBboxPatch((inner_x0, cursor_y - quote_h), inner_x1 - inner_x0, quote_h,
                                          boxstyle="round,pad=0,rounding_size=0.008",
                                          facecolor="none", edgecolor=COVER_GOLD, linewidth=0.9, zorder=3))
    q_lines = _wrap_lines(abertura, width=74, max_lines=2)
    for i, ln in enumerate(q_lines):
        ax.text((inner_x0 + inner_x1) / 2, cursor_y - quote_h / 2 + (len(q_lines) - 1 - 2 * i) * 0.012,
                ln, ha="center", va="center", fontsize=9.3, color=CARD_INK,
                family="serif", style="italic", zorder=4)
    cursor_y -= quote_h + 0.012

    band_top = cursor_y
    band_bottom = CARD_BOTTOM + pad
    left_w = 0.415
    left_x0, left_x1 = inner_x0, inner_x0 + (inner_x1 - inner_x0) * left_w
    right_x0 = left_x1 + 0.012
    right_x1 = inner_x1

    # ---- coluna esquerda: 3 blocos com icone (personalidade / relacionamentos / crescimento)
    boxes = [
        ("✦", "#3f6fb0", "PERSONALIDADE",
         sections.get("capa_personalidade") or (sections.get("personalidade_subtitle") or
             f"Sol em {sol['sign']}, Lua em {lua['sign']}: uma combinação própria de identidade e sensibilidade." if sol and lua else "")),
        ("♥", "#c0567a", "RELACIONAMENTOS",
         sections.get("capa_relacionamentos") or (sections.get("afeto_subtitle") or
             "Você valoriza vínculos verdadeiros, com espaço para individualidade e crescimento em conjunto.")),
        ("▲", "#3f9e7a", "CRESCIMENTO",
         sections.get("capa_crescimento") or
             "Desenvolver constância e presença ajuda a transformar grandes visões em realizações duradouras."),
    ]
    box_h = (band_top - band_bottom) / 3.0
    for i, (glyph, color, title, text) in enumerate(boxes):
        by1 = band_top - i * box_h
        by0 = by1 - box_h + 0.006
        icon_cx, icon_cy = left_x0 + 0.017, by1 - box_h * 0.30
        ax.add_patch(mpatches.Circle((icon_cx, icon_cy), 0.016, facecolor=color, alpha=0.85, edgecolor="none", zorder=4))
        ax.text(icon_cx, icon_cy, glyph, ha="center", va="center", fontsize=8.5, color="#ffffff",
                family="DejaVu Sans", zorder=5)
        text_x = icon_cx + 0.032
        ax.text(text_x, by1 - box_h * 0.14, title, ha="left", va="center", fontsize=8.3,
                color=CARD_INK, fontweight="bold", family=FONT, zorder=4)
        lines = _wrap_lines(text, width=46, max_lines=3)
        for li, ln in enumerate(lines):
            ax.text(text_x, by1 - box_h * 0.14 - 0.017 - li * 0.0135, ln, ha="left", va="center",
                    fontsize=6.9, color=CARD_SUB, family=FONT, zorder=4)

    # ---- coluna direita: tabela de posicoes planetarias + dados de nascimento
    table_w = (right_x1 - right_x0) * 0.60
    tb_x0, tb_x1 = right_x0, right_x0 + table_w
    db_x0, db_x1 = tb_x0 + table_w + 0.012, right_x1

    table_planets = [p for p in planets if p["name"] not in ("Ascendente", "Meio do Céu")]
    title_gap = 0.020
    hdr_h = 0.011
    band_h = band_top - band_bottom
    row_h = max(0.0088, min(0.0158, (band_h - title_gap - hdr_h) / max(1, len(table_planets))))
    row_font = 5.9 if row_h >= 0.0135 else (5.4 if row_h >= 0.0110 else 4.9)

    ax.text((tb_x0 + tb_x1) / 2, band_top - 0.010, _track("POSIÇÕES PLANETÁRIAS E CASAS"),
            ha="center", va="center", fontsize=6.8, color=COVER_NAVY, fontweight="bold", family=FONT, zorder=4)
    table_rows_y0 = band_top - title_gap
    col_x = [tb_x0 + 0.004, tb_x0 + (tb_x1 - tb_x0) * 0.34, tb_x0 + (tb_x1 - tb_x0) * 0.62, tb_x0 + (tb_x1 - tb_x0) * 0.87]
    ax.add_patch(mpatches.Rectangle((tb_x0, table_rows_y0 - hdr_h), tb_x1 - tb_x0, hdr_h,
                                     facecolor=COVER_NAVY, edgecolor="none", zorder=4))
    for cx, htxt in zip(col_x, ("PLANETA", "SIGNO", "GRAU", "CASA")):
        ax.text(cx, table_rows_y0 - hdr_h / 2, htxt, ha="left", va="center", fontsize=min(5.6, row_font + 0.3),
                color="#ffffff", fontweight="bold", family=FONT, zorder=5)

    y = table_rows_y0 - hdr_h
    for i, p in enumerate(table_planets):
        y -= row_h
        if i % 2 == 1:
            ax.add_patch(mpatches.Rectangle((tb_x0, y), tb_x1 - tb_x0, row_h,
                                             facecolor="#efe9d8", edgecolor="none", zorder=3.5))
        house_txt = re.sub(r"[^0-9]", "", str(p.get("house", "-"))) or "-"
        vals = (p["name"], p["sign"], f'{p["deg"]:02d}°{p["min"]:02d}’', house_txt)
        for cx, v in zip(col_x, vals):
            ax.text(cx, y + row_h / 2, v, ha="left", va="center", fontsize=row_font,
                    color=CARD_INK, family=FONT, zorder=4)

    # dados de nascimento
    ax.text((db_x0 + db_x1) / 2, band_top - 0.010, "DADOS DE NASCIMENTO", ha="center", va="center",
            fontsize=6.3, color=COVER_NAVY, fontweight="bold", family=FONT, zorder=4)
    db_rows = [
        ("Data:", _fmt_date_short(birth.get("date", ""))),
        ("Hora:", f'{birth.get("time", "")}h'),
        ("Local:", birth.get("place", "")),
        ("Lat.:", _dm(birth["lat"], "S", "N") if birth.get("lat") is not None else "-"),
        ("Long.:", _dm(birth["lon"], "L", "O") if birth.get("lon") is not None else "-"),
        ("Sistema:", "Tropical"),
        ("Casas:", "Placidus"),
    ]
    db_row_h = min(0.0165, max(0.011, row_h * 1.05))
    dy = band_top - 0.028
    for label, value in db_rows:
        ax.text(db_x0 + 0.004, dy, label, ha="left", va="center", fontsize=5.7,
                color=CARD_SUB, family=FONT, zorder=4)
        ax.text(db_x1 - 0.004, dy, value, ha="right", va="center", fontsize=5.7,
                color=CARD_INK, fontweight="bold", family=FONT, zorder=4, wrap=True)
        dy -= db_row_h
    compass_cy = max(band_bottom + 0.030, dy - 0.010)
    ax.add_patch(mpatches.Circle((db_x1 - 0.028, compass_cy), 0.020,
                                  facecolor="none", edgecolor=COVER_GOLD, linewidth=0.7, zorder=4))
    for ang, lbl in ((90, "N"), (270, "S"), (0, "L"), (180, "O")):
        rad = math.radians(ang)
        ax.text(db_x1 - 0.028 + 0.026 * math.cos(rad), compass_cy + 0.026 * math.sin(rad),
                lbl, ha="center", va="center", fontsize=4.6, color=COVER_GOLD, family=FONT, zorder=4)
    ax.text(db_x1 - 0.028, compass_cy, "✦", ha="center", va="center", fontsize=6,
            color=COVER_GOLD, zorder=4)

    # -------------------------------------------------- linha asc/mc + fecho
    asc_label = birth.get("ascendant_label") or f'{ascendant["deg"]:02d}°{ascendant["min"]:02d}’ {ascendant["sign"]}'
    mc_label = birth.get("midheaven_label") or (f'{mc["deg"]:02d}°{mc["min"]:02d}’ {mc["sign"]}' if mc else "-")
    ax.text(0.5, 0.132, f'ASCENDENTE: {ascendant["sign"].upper()} {asc_label.split()[0]}   •   '
            f'MEIO DO CÉU: {(mc["sign"].upper() if mc else "")} {mc_label.split()[0]}',
            ha="center", va="center", fontsize=8, color=COVER_GOLD, family=FONT, zorder=3)

    fechamento = sections.get("capa_fechamento") or (
        f"Sua mente cria possibilidades,\nseu coração escolhe o que vale a pena realizar."
        if sol and sol["sign"] in ("Gêmeos", "Virgem") else
        f'"{first_name} une o que sente e o que constrói em uma única direção."')
    ax.plot([0.34, 0.42], [0.090, 0.090], color=COVER_GOLD, linewidth=0.6, alpha=0.7, zorder=3)
    ax.text(0.5, 0.090, "✦", ha="center", va="center", fontsize=7, color=COVER_GOLD, zorder=3)
    ax.plot([0.58, 0.66], [0.090, 0.090], color=COVER_GOLD, linewidth=0.6, alpha=0.7, zorder=3)
    fech_lines = (fechamento or "").split("\n") if "\n" in (fechamento or "") else _wrap_lines(fechamento, width=52, max_lines=2)
    for i, ln in enumerate(fech_lines):
        ax.text(0.5, 0.062 - i * 0.026, ln, ha="center", va="center", fontsize=12.5,
                color=COVER_GOLD_LIGHT, family="serif", style="italic", zorder=3)

    fig.savefig(out_path, facecolor=fig.get_facecolor())
    plt.close(fig)
    return out_path

def _mini_pos(chart, planet_name):
    p = next((x for x in chart["planets"] if x["name"] == planet_name), None)
    return p["sign"] if p else "-"

def _mini_triad_line(chart):
    sol, lua = _mini_pos(chart, "Sol"), _mini_pos(chart, "Lua")
    asc = chart["ascendant"]["sign"]
    return f"Sol {sol} • Lua {lua} • Asc {asc}"

def build_pet_synastry_cover_png(data, out_path, owner_photo_path=None, pet_photo_path=None):
    """Capa navy/dourada com dois sujeitos (tutor + pet), reaproveitando as
    primitivas de `build_placeholder_cover_png` (fundo estrelado, paleta,
    `_photo_rgba_cover`, `_draw_wheel`) num layout de duas colunas: uma para o
    tutor, outra para o pet. Cada coluna mostra a foto enviada (se houver) ou,
    na ausencia dela, a roda natal daquele sujeito."""
    owner, pet = data["owner"], data["pet"]
    owner_name, pet_name = owner["name"], pet["name"]

    fig_w, fig_h = 8.27, 11.69  # A4 portrait
    dpi = 220
    fig = plt.figure(figsize=(fig_w, fig_h), dpi=dpi)
    fig.patch.set_facecolor("#0a213b")

    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.set_facecolor("#0a213b")

    rng = __import__("random").Random(hash(owner_name + pet_name) % (2**31))
    for _ in range(110):
        sx, sy = rng.uniform(0.03, 0.97), rng.uniform(0.03, 0.97)
        ax.plot(sx, sy, marker="*", markersize=rng.uniform(1.2, 3.2),
                color="#c9a24b", alpha=rng.uniform(0.15, 0.55), zorder=1)

    ax.text(0.5, 0.955, "SINASTRIA", ha="center", fontsize=27, color="#e9c98a",
            fontweight="bold", family="serif", zorder=3)
    ax.text(0.5, 0.920, "COM PET", ha="center", fontsize=18, color="#e9c98a",
            fontweight="bold", family="serif", zorder=3)
    ax.plot([0.32, 0.68], [0.900, 0.900], color="#b79242", linewidth=0.8, alpha=0.7, zorder=3)

    ax.text(0.5, 0.868, f"{owner_name.upper()}  &  {pet_name.upper()}", ha="center", fontsize=14,
            color="#ffffff", family="serif", zorder=3)
    ax.text(0.5, 0.842, _format_birth_line(owner["birth"]), ha="center", fontsize=9,
            color="#cbd3dc", family="serif", zorder=3)
    ax.text(0.5, 0.822, _format_birth_line(pet["birth"]), ha="center", fontsize=9,
            color="#cbd3dc", family="serif", zorder=3)

    col_boxes = [(0.035, 0.315, 0.435, 0.455), (0.53, 0.315, 0.435, 0.455)]
    names = [owner_name, pet_name]
    photos = [owner_photo_path, pet_photo_path]
    charts = [owner, pet]

    for (bx0, by0, bw, bh), nm, photo_path, ch in zip(col_boxes, names, photos, charts):
        if photo_path:
            try:
                rgba = _photo_rgba_cover(photo_path, bw * fig_w, bh * fig_h, dpi, fade_bottom=0.30)
                ax.imshow(rgba, extent=[bx0, bx0 + bw, by0, by0 + bh], zorder=2, aspect="auto")
            except Exception:
                photo_path = None
        if not photo_path:
            wheel_h = bw * fig_w / fig_h
            wheel_y = by0 + (bh - wheel_h) / 2
            wheel_ax = fig.add_axes([bx0, wheel_y, bw, wheel_h])
            wheel_ax.set_xlim(-1.18, 1.18); wheel_ax.set_ylim(-1.18, 1.18)
            wheel_ax.set_aspect("equal"); wheel_ax.axis("off"); wheel_ax.patch.set_alpha(0)
            _draw_wheel(wheel_ax, ch["planets"], ch["houses"], ch["ascendant"], bg=None, dot_r=0.036)

        cx = bx0 + bw / 2
        ax.text(cx, by0 - 0.018, nm.upper(), ha="center", fontsize=10, color="#ffffff",
                family="serif", zorder=3)
        ax.text(cx, by0 - 0.036, _mini_triad_line(ch), ha="center", fontsize=7.6,
                color="#cbd3dc", family="serif", style="italic", zorder=3)

    ax.text(0.5, 0.248, "TRÍADE COMPARADA", ha="center", fontsize=9.5, color="#e9c98a",
            family="serif", fontweight="bold", zorder=3)
    col_x = [0.15, 0.45, 0.75]
    ax.text(col_x[1], 0.228, owner_name.split()[0].upper(), ha="center", fontsize=7.6,
            color="#b79242", family="serif", zorder=3)
    ax.text(col_x[2], 0.228, pet_name.upper(), ha="center", fontsize=7.6,
            color="#b79242", family="serif", zorder=3)
    rows = [
        ("Sol", _mini_pos(owner, "Sol"), _mini_pos(pet, "Sol")),
        ("Lua", _mini_pos(owner, "Lua"), _mini_pos(pet, "Lua")),
        ("Ascendente", owner["ascendant"]["sign"], pet["ascendant"]["sign"]),
    ]
    for i, (label, ov, pv) in enumerate(rows):
        y = 0.206 - i * 0.019
        ax.text(col_x[0], y, label, ha="left", fontsize=8.6, color="#ffffff", family="serif", zorder=3)
        ax.text(col_x[1], y, ov, ha="center", fontsize=8.6, color="#cbd3dc", family="serif", zorder=3)
        ax.text(col_x[2], y, pv, ha="center", fontsize=8.6, color="#cbd3dc", family="serif", zorder=3)

    ax.plot([0.32, 0.68], [0.095, 0.095], color="#b79242", linewidth=0.8, alpha=0.7, zorder=3)
    ax.text(0.5, 0.065, '"Entre vocês, o vínculo fala uma língua que dispensa palavras."',
            ha="center", va="center", fontsize=10, color="#e9c98a", style="italic",
            family="serif", zorder=3, linespacing=1.6)

    fig.savefig(out_path, facecolor=fig.get_facecolor())
    plt.close(fig)
    return out_path

# -------------------------------------------------------------- page manifest

def build_pages(data):
    """Constroi a lista de 24 paginas de conteudo (sem a capa) a partir de `data`
    (ver SCHEMA.md). Campos ausentes usam texto de fallback minimo."""
    name = data["name"]
    s = data.get("sections", {})
    birth = data["birth"]
    planets = data["planets"]
    houses = data["houses"]

    def g(key, default=""):
        return s.get(key, default)

    pages = []

    pages.append(dict(eyebrow="LEITURA DO MAPA", title="ANTES DE COMEÇAR",
        subtitle=g("intro_subtitle", "Uma leitura simbólica para orientar reflexão, não um roteiro fechado sobre o futuro."),
        blocks=[
            ('callout', "Uma observação essencial", g("intro_observacao",
                "Este relatório utiliza a linguagem simbólica da astrologia ocidental tropical, com sistema de casas Placidus. "
                "As interpretações não constituem previsão científica nem determinam escolhas, acontecimentos ou capacidades.")),
            ('p', g("intro_p1", "")),
            ('p', g("intro_p2", "")),
            ('callout', "Como aproveitar melhor este relatório", g("intro_como_aproveitar",
                "Observe padrões recorrentes, compare-os com sua experiência e use a previsão anual como um convite à consciência.")),
        ]))

    pages.append(dict(eyebrow="1. DADOS DE NASCIMENTO", title="DADOS DE NASCIMENTO",
        subtitle="Pontos técnicos que organizam a leitura da carta.",
        blocks=[
            ('kv', [
                ("Nome completo", birth.get("full_name", name)),
                ("Data de nascimento", birth.get("date", "")),
                ("Horário informado", birth.get("time", "")),
                ("Local", birth.get("place", "")),
                ("Sistema zodiacal", birth.get("zodiac_system", "Astrologia ocidental tropical")),
                ("Sistema de casas", birth.get("house_system", "Placidus")),
                ("Ascendente calculado", birth.get("ascendant_label", "")),
                ("Meio do Céu calculado", birth.get("midheaven_label", "")),
            ]),
            ('p', g("dados_nascimento_texto",
                "O Ascendente e o Meio do Céu calculados acima organizam a leitura de todo o restante do "
                "mapa nas próximas páginas — o primeiro marca a porta de entrada da sua energia no mundo, "
                "o segundo aponta a direção da sua vocação pública.")),
        ]))

    pages.append(dict(eyebrow="MAPA TÉCNICO", title="RODA NATAL TÉCNICA",
        subtitle="Visualização das casas, signos e principais planetas do nascimento.",
        blocks=[
            ('image', data["_wheel_png_path"], 10.5),
            ('p', g("mapa_tecnico_texto",
                "A roda natal organiza os doze signos no anel externo e as doze casas no anel interno. "
                "Os pontos coloridos indicam a posição simbólica de cada planeta.")),
        ]))

    pages.append(dict(eyebrow="2. SOL, LUA E ASCENDENTE", title="TRÍADE PRINCIPAL",
        subtitle="A estrutura que une identidade, emoção e presença.",
        blocks=[
            ('subp', g("triade_items", [])),
            ('callout', "Síntese da tríade", g("triade_sintese", "")),
        ]))

    planet_rows = [(p["name"], f'{p["deg"]:02d}°{p["min"]:02d} {p["sign"]}', p.get("house", "-"), p.get("key", "")) for p in planets]
    pages.append(dict(eyebrow="3. POSIÇÕES PLANETÁRIAS", title="POSIÇÕES PLANETÁRIAS E CASAS",
        subtitle="Os principais pontos do mapa natal e suas áreas de manifestação.",
        blocks=[
            ('table', ["Ponto", "Posição", "Casa", "Chave simbólica"], planet_rows),
            ('p', g("posicoes_texto", "")),
        ]))

    house_lines = []
    hs = {h["n"]: h for h in houses}
    for a, b in [(1, 7), (2, 8), (3, 9), (4, 10), (5, 11), (6, 12)]:
        ha, hb = hs.get(a), hs.get(b)
        if ha and hb:
            house_lines.append(f'Casa {a}: {ha["deg"]}°{ha["min"]:02d} {ha["sign"]}      Casa {b}: {hb["deg"]}°{hb["min"]:02d} {hb["sign"]}')
    pages.append(dict(eyebrow="CASAS ASTROLÓGICAS", title="CASAS E EIXOS DO MAPA",
        subtitle="As áreas da vida que recebem maior densidade simbólica.",
        blocks=[
            ('p', g("casas_texto", "")),
            ('linelist', house_lines),
            ('callout', g("eixo1_title", "Eixo Ascendente–Descendente"), g("eixo1_text", "")),
            ('callout', g("eixo2_title", "Eixo Fundo do Céu–Meio do Céu"), g("eixo2_text", "")),
        ]))

    pages.append(dict(eyebrow="4. PERSONALIDADE", title="PERSONALIDADE E TRAÇOS CENTRAIS",
        subtitle=g("personalidade_subtitle", "Traços centrais da personalidade."),
        blocks=[('p', t) for t in g("personalidade", [])] +
               [('bullets', "Síntese prática", g("personalidade_bullets", []))]))

    pages.append(dict(eyebrow="5. EMOÇÕES E RELAÇÕES", title="MUNDO EMOCIONAL E RELACIONAMENTOS",
        subtitle=g("emocional_subtitle", "O mundo emocional e os vínculos afetivos."),
        blocks=[('p', t) for t in g("emocional", [])] +
               [('bullets', "Síntese prática", g("emocional_bullets", []))]))

    pages.append(dict(eyebrow="6. COMUNICAÇÃO E MENTE", title="COMUNICAÇÃO, PENSAMENTO E DECISÕES",
        subtitle=g("comunicacao_subtitle", "Como a mente processa, decide e se expressa."),
        blocks=[('p', t) for t in g("comunicacao", [])] +
               [('bullets', "Síntese prática", g("comunicacao_bullets", []))]))

    pages.append(dict(eyebrow="7. AFETO E VALORES", title="AMOR, AFETO, DESEJO E VALORES",
        subtitle=g("afeto_subtitle", "Amor, afeto, desejo e valores pessoais."),
        blocks=[('p', t) for t in g("afeto", [])] +
               [('bullets', "Síntese prática", g("afeto_bullets", []))]))

    pages.append(dict(eyebrow="8. CARREIRA E PROPÓSITO", title="CARREIRA, IMAGEM PÚBLICA E PROPÓSITO",
        subtitle=g("carreira_subtitle", "Carreira, imagem pública e propósito de vida."),
        blocks=[('p', t) for t in g("carreira", [])] +
               [('bullets', "Síntese prática", g("carreira_bullets", []))]))

    pages.append(dict(eyebrow="9. TALENTOS", title="FORÇAS E TALENTOS NATURAIS",
        subtitle="O que ganha potência quando você confia na própria natureza.",
        blocks=[('list', [(t["lead"], t["text"]) for t in g("talentos", [])]),
                ('callout', "Talento central", g("talentos_callout", ""))]))

    pages.append(dict(eyebrow="10. DESENVOLVIMENTO PESSOAL", title="OPORTUNIDADES DE CRESCIMENTO",
        subtitle="Pontos que podem se tornar mais livres e conscientes.",
        blocks=[('numlist', [(i + 1, it["title"], it["text"]) for i, it in enumerate(g("crescimento", []))])]))

    pages.append(dict(eyebrow="11. ASPECTOS PRINCIPAIS - I", title="ASPECTOS ASTROLÓGICOS EM DESTAQUE",
        subtitle="Diálogos simbólicos que se combinam com naturalidade.",
        blocks=[('subp', g("aspectos_harmonicos_items", []))]))

    pages.append(dict(eyebrow="11. ASPECTOS PRINCIPAIS - II", title="ASPECTOS ASTROLÓGICOS EM DESTAQUE",
        subtitle="Tensões que pedem consciência, escolha e refinamento.",
        blocks=[('subp', g("aspectos_tensos_items", []))]))

    pages.append(dict(eyebrow="12. CICLOS E TRÂNSITOS", title="PREVISÃO SIMBÓLICA PARA OS PRÓXIMOS 12 MESES",
        subtitle=g("previsao_subtitle", ""),
        blocks=[('p', t) for t in g("previsao_overview", [])]))

    for q in g("trimestres", []):
        pages.append(dict(eyebrow=q["eyebrow"], title="PREVISÃO SIMBÓLICA", subtitle=q["subtitle"], blocks=[
            ('p', q["p1"]), ('p', q["p2"]),
            ('callout', q["opp_title"], q["opp_text"]),
            ('callout', q["dec_title"], q["dec_text"]),
        ]))

    pages.append(dict(eyebrow="13. SÍNTESE FINAL", title="RESUMO FINAL",
        subtitle="A essência que se destaca quando se observa o mapa como conjunto.",
        blocks=[('p', t) for t in g("sintese_final", [])] +
               [('quote', g("sintese_quote", ""))]))

    if houses:
        cusp_rows = [(f'Casa {h["n"]}', f'{h["deg"]}°{h["min"]:02d} {h["sign"]}') for h in sorted(houses, key=lambda x: x["n"])]
        pages.append(dict(eyebrow="DADOS COMPLEMENTARES", title="APÊNDICE TÉCNICO",
            subtitle="Cúspides das casas calculadas para esta leitura.",
            blocks=[('p', g("apendice_texto", "")), ('table', ["Casa", "Cúspide"], cusp_rows)]))

    all_asp = g("aspectos_harmonicos", []) + g("aspectos_tensos", [])
    if all_asp:
        pages.append(dict(eyebrow="DADOS COMPLEMENTARES", title="ASPECTOS SELECIONADOS",
            subtitle="Relações com orbes aproximados.",
            blocks=[('table', ["Aspecto", "Orbe aproximado"], [(a["aspecto"], a["orbe"]) for a in all_asp])]))

    pages.append(dict(eyebrow="NOTA FINAL", title="FECHO",
        subtitle="Um mapa não encerra uma pessoa. Ele abre uma conversa.",
        blocks=[('p', t) for t in g("nota_final", [])] +
               [('quote', g("nota_quote", g("sintese_quote", "")))]))

    return pages

_ELEMENT_OF_SIGN = {
    "Áries": "fogo", "Leão": "fogo", "Sagitário": "fogo",
    "Touro": "terra", "Virgem": "terra", "Capricórnio": "terra",
    "Gêmeos": "ar", "Libra": "ar", "Aquário": "ar",
    "Câncer": "água", "Escorpião": "água", "Peixes": "água",
}

def _chart_shape_caption(chart_planets):
    """Legenda curta e nao-generica para a pagina da roda natal tecnica,
    citando o elemento predominante calculado a partir dos planetas do mapa
    (em vez de uma frase estatica identica em todos os relatorios)."""
    counts = {}
    for p in chart_planets:
        el = _ELEMENT_OF_SIGN.get(p.get("sign"))
        if el:
            counts[el] = counts.get(el, 0) + 1
    if not counts:
        return ("A roda natal organiza os doze signos no anel externo e as doze casas no anel interno. "
                "Os pontos coloridos indicam a posição simbólica de cada planeta.")
    top_el, top_n = max(counts.items(), key=lambda x: x[1])
    return (f"A roda natal organiza os doze signos no anel externo e as doze casas no anel interno. "
            f"Os pontos coloridos indicam a posição simbólica de cada planeta — aqui, com predomínio do "
            f"elemento {top_el} ({top_n} dos pontos calculados).")

def build_pet_synastry_pages(data):
    """Constroi a lista de paginas de conteudo (sem a capa) para o relatorio
    de sinastria com pet, a partir de `data`:
    {"owner": {...chart do tutor...}, "pet": {...chart do pet + breed/color/time_estimated...},
     "cross_aspects": [...], "house_overlay_owner": [...], "house_overlay_pet": [...],
     "sections": {...}, "_wheel_png_owner": path, "_wheel_png_pet": path}."""
    owner, pet = data["owner"], data["pet"]
    s = data.get("sections", {})
    owner_name, pet_name = owner["name"], pet["name"]

    def g(key, default=""):
        return s.get(key, default)

    pages = []

    intro_blocks = [
        ('callout', "Uma observação essencial", g("intro_observacao",
            "Este relatório cruza a astrologia ocidental tropical, com sistema de casas Placidus, do tutor "
            "com uma leitura simbólica do temperamento do pet. As interpretações não constituem previsão "
            "científica nem determinam comportamento, escolhas ou saúde.")),
    ]
    disclaimer = g("time_estimated_disclaimer", "")
    if disclaimer:
        intro_blocks.append(('callout', "Sobre o horário de nascimento do pet", disclaimer))
    intro_blocks += [
        ('p', g("intro_apresentacao", "")),
        ('callout', "Como aproveitar melhor este relatório", g("intro_como_aproveitar",
            "Observe padrões recorrentes na convivência entre vocês e use este relatório como um convite à consciência.")),
    ]
    pages.append(dict(eyebrow="LEITURA DA SINASTRIA", title="ANTES DE COMEÇAR",
        subtitle=g("intro_subtitle", f"Uma leitura simbólica do vínculo entre {owner_name} e {pet_name}."),
        blocks=intro_blocks))

    pages.append(dict(eyebrow="1. DADOS DE NASCIMENTO", title="DADOS DE NASCIMENTO",
        subtitle="Pontos técnicos que organizam a leitura dos dois mapas.",
        blocks=[
            ('p', f"**{owner_name}** (tutor)"),
            ('kv', [
                ("Nome completo", owner["birth"].get("full_name", owner_name)),
                ("Data de nascimento", owner["birth"].get("date", "")),
                ("Horário informado", owner["birth"].get("time", "")),
                ("Local", owner["birth"].get("place", "")),
                ("Ascendente calculado", owner["birth"].get("ascendant_label", "")),
            ]),
            ('p', f"**{pet_name}** (pet — {pet.get('breed','')}, {pet.get('color','')})"),
            ('kv', [
                ("Nome", pet["birth"].get("full_name", pet_name)),
                ("Data de nascimento", pet["birth"].get("date", "")),
                ("Horário " + ("estimado (meio-dia)" if pet.get("time_estimated") else "informado"), pet["birth"].get("time", "")),
                ("Local", pet["birth"].get("place", "")),
                ("Ascendente calculado", pet["birth"].get("ascendant_label", "")),
            ]),
            ('p', g("dados_nascimento_texto",
                f"O Ascendente de {owner_name} organiza a leitura da tríade do tutor nas próximas páginas, "
                f"e o Ascendente calculado para {pet_name} situa a presença do pet no ambiente — a base "
                f"técnica sobre a qual a leitura cruzada do vínculo entre os dois se apoia.")),
        ]))

    pages.append(dict(eyebrow="MAPA TÉCNICO", title=f"RODA NATAL — {owner_name.upper()}",
        subtitle="Visualização das casas, signos e principais planetas do tutor.",
        blocks=[
            ('image', data["_wheel_png_owner"], 10.5),
            ('p', _chart_shape_caption(owner["planets"])),
        ]))

    pages.append(dict(eyebrow="MAPA TÉCNICO", title=f"RODA NATAL — {pet_name.upper()}",
        subtitle="Visualização das casas, signos e principais planetas do pet.",
        blocks=[
            ('image', data["_wheel_png_pet"], 10.5),
            ('p', _chart_shape_caption(pet["planets"]) + " A mesma leitura técnica se aplica ao mapa do "
                  "pet, ainda que a interpretação a seguir respeite sua natureza animal."),
        ]))

    pages.append(dict(eyebrow="2. TRÍADE DO TUTOR", title=f"TRÍADE DE {owner_name.upper()}",
        subtitle="A estrutura que une identidade, emoção e presença do tutor.",
        blocks=[
            ('p', g("owner_sol", "")),
            ('p', g("owner_lua", "")),
            ('p', g("owner_ascendente", "")),
            ('callout', "Síntese da tríade", g("owner_sintese", "")),
        ]))

    pages.append(dict(eyebrow="3. ESSÊNCIA DO PET", title=f"ESSÊNCIA DE {pet_name.upper()}",
        subtitle="Temperamento e presença do pet, em registro simbólico.",
        blocks=[
            ('p', g("pet_sol", "")),
            ('p', g("pet_lua", "")),
            ('p', g("pet_ascendente", "")),
        ]))

    pet_planet_rows = [(p["name"], f'{p["deg"]:02d}°{p["min"]:02d} {p["sign"]}', p.get("house", "-"), p.get("key", ""))
                        for p in pet["planets"]]
    pages.append(dict(eyebrow="4. POSIÇÕES DO PET", title="POSIÇÕES E TEMPERAMENTO",
        subtitle=f"Os principais pontos do mapa de {pet_name} e suas áreas de manifestação.",
        blocks=[
            ('table', ["Ponto", "Posição", "Casa", "Chave simbólica"], pet_planet_rows),
            ('p', g("pet_posicoes_texto", "")),
        ]))

    pages.append(dict(eyebrow="5. O VÍNCULO", title="O QUE CONECTA VOCÊS",
        subtitle=f"A leitura cruzada entre {owner_name} e {pet_name}.",
        blocks=[
            ('p', g("vinculo_texto", "")),
            ('callout', "Essência do vínculo", g("vinculo_callout", "")),
        ]))

    pages.append(dict(eyebrow="6. ASPECTOS CRUZADOS - I", title="ASPECTOS CRUZADOS EM DESTAQUE",
        subtitle="Diálogos simbólicos que se combinam com naturalidade.",
        blocks=[
            ('kv', [(a["aspecto"], f'Orbe aproximado: {a["orbe"]}') for a in g("aspectos_cruzados_harmonicos", [])]),
            ('p', g("aspectos_cruzados_harmonicos_texto", "")),
        ]))

    pages.append(dict(eyebrow="6. ASPECTOS CRUZADOS - II", title="ASPECTOS CRUZADOS EM DESTAQUE",
        subtitle="Tensões que pedem consciência, escolha e refinamento.",
        blocks=[
            ('kv', [(a["aspecto"], f'Orbe aproximado: {a["orbe"]}') for a in g("aspectos_cruzados_tensos", [])]),
            ('p', g("aspectos_cruzados_tensos_texto", "")),
        ]))

    overlay_lines = (
        [f'{o["planet"]} de {owner_name} na casa {o["house_in_b"]} de {pet_name}'
         for o in data.get("house_overlay_owner", [])[:6]] +
        [f'{o["planet"]} de {pet_name} na casa {o["house_in_b"]} de {owner_name}'
         for o in data.get("house_overlay_pet", [])[:6]]
    )
    pages.append(dict(eyebrow="7. CASAS CRUZADAS", title="SOBREPOSIÇÃO DE CASAS",
        subtitle="Em quais áreas de vida um ativa o outro.",
        blocks=[
            ('linelist', overlay_lines),
            ('p', g("casas_cruzadas_texto", "")),
        ]))

    pages.append(dict(eyebrow="8. ROTINA A DOIS", title="ROTINA E COTIDIANO",
        subtitle="Como o vínculo aparece no dia a dia.",
        blocks=[('p', g("rotina_texto", ""))]))

    pages.append(dict(eyebrow="9. DESAFIOS", title="DESAFIOS E CRESCIMENTO",
        subtitle="Pontos de atenção que fortalecem a dupla quando trabalhados.",
        blocks=[('numlist', [(i + 1, it["title"], it["text"]) for i, it in enumerate(g("desafios", []))])]))

    pages.append(dict(eyebrow="10. CICLOS E PREVISÃO", title="PREVISÃO SIMBÓLICA COMPARTILHADA",
        subtitle="Os próximos ciclos para o vínculo entre vocês.",
        blocks=[('p', g("previsao_texto", ""))]))

    pages.append(dict(eyebrow="11. SÍNTESE FINAL", title="RESUMO FINAL",
        subtitle="A essência que se destaca quando se observa a dupla como conjunto.",
        blocks=[('p', t) for t in g("sintese_final", [])] +
               [('quote', g("sintese_quote", ""))]))

    if owner["houses"]:
        cusp_rows = [(f'Casa {h["n"]}', f'{h["deg"]}°{h["min"]:02d} {h["sign"]}') for h in sorted(owner["houses"], key=lambda x: x["n"])]
        pages.append(dict(eyebrow="DADOS COMPLEMENTARES", title=f"APÊNDICE TÉCNICO — {owner_name.upper()}",
            subtitle="Cúspides das casas calculadas para o tutor.",
            blocks=[('table', ["Casa", "Cúspide"], cusp_rows)]))

    if pet["houses"]:
        cusp_rows = [(f'Casa {h["n"]}', f'{h["deg"]}°{h["min"]:02d} {h["sign"]}') for h in sorted(pet["houses"], key=lambda x: x["n"])]
        pages.append(dict(eyebrow="DADOS COMPLEMENTARES", title=f"APÊNDICE TÉCNICO — {pet_name.upper()}",
            subtitle="Cúspides das casas calculadas para o pet.",
            blocks=[('table', ["Casa", "Cúspide"], cusp_rows)]))

    all_cross = g("aspectos_cruzados_harmonicos", []) + g("aspectos_cruzados_tensos", [])
    if all_cross:
        pages.append(dict(eyebrow="DADOS COMPLEMENTARES", title="ASPECTOS CRUZADOS SELECIONADOS",
            subtitle="Relações entre os dois mapas com orbes aproximados.",
            blocks=[('table', ["Aspecto", "Orbe aproximado"], [(a["aspecto"], a["orbe"]) for a in all_cross])]))

    pages.append(dict(eyebrow="NOTA FINAL", title="FECHO",
        subtitle="Um mapa não encerra uma relação. Ele abre uma conversa.",
        blocks=[('p', t) for t in g("nota_final", [])]))

    return pages

# ------------------------------------------------------------------ document

def build_docx_bytes(data, cover_image_path=None):
    """data: dict seguindo SCHEMA.md. cover_image_path: caminho para JPG/PNG de
    capa (retrato ja gerado, se existir); se None, uma capa navy/dourada simples
    e desenhada automaticamente."""
    tmp_dir = tempfile.mkdtemp(prefix="report_")
    wheel_png = os.path.join(tmp_dir, "wheel.png")
    build_wheel_png(data["planets"], data["houses"], data["ascendant"], wheel_png)
    data = dict(data)
    data["_wheel_png_path"] = wheel_png

    uploaded_photo_path = cover_image_path
    cover_image_path = os.path.join(tmp_dir, "cover.png")
    build_placeholder_cover_png(data["name"], data["birth"], data["planets"],
                                 data["houses"], data["ascendant"], cover_image_path,
                                 photo_path=uploaded_photo_path, sections=data.get("sections"),
                                 aspects=data.get("aspects"))

    name = data["name"]
    running = f"MAPA ASTRAL PERSONALIZADO • {name.upper()}"

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(9.2); normal.font.color.rgb = BODY_INK

    sec0 = doc.sections[0]
    sec0.page_width = Cm(21.0); sec0.page_height = Cm(29.7)
    sec0.left_margin = sec0.right_margin = sec0.top_margin = sec0.bottom_margin = Cm(0)
    sec0.header_distance = sec0.footer_distance = Cm(0)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.add_run().add_picture(cover_image_path, width=Cm(21.0), height=Cm(29.7))

    sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec1.page_width = Cm(21.0); sec1.page_height = Cm(29.7)
    sec1.left_margin = sec1.right_margin = Cm(2.35)
    sec1.top_margin = sec1.bottom_margin = Cm(1.55)
    sec1.header_distance = Cm(0.5); sec1.footer_distance = Cm(0.6)
    sec1.header.is_linked_to_previous = False
    sec1.footer.is_linked_to_previous = False

    hp = sec1.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run(running); hr.font.name = FONT; hr.font.size = Pt(6.5); hr.font.bold = True; hr.font.color.rgb = GOLD
    para_bottom_border(hp, color=GOLD_LINE, sz=5)

    fp = sec1.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(running + " • "); fr.font.name = FONT; fr.font.size = Pt(6.3); fr.font.color.rgb = FOOTER_INK
    add_page_field(fp)

    pages = build_pages(data)
    for i, pg in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        render_head(doc, pg['eyebrow'], pg['title'], pg['subtitle'])
        render_blocks(doc, pg['blocks'])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def build_pet_synastry_docx_bytes(data, owner_cover_path=None, pet_cover_path=None):
    """data: {"owner":{...}, "pet":{...}, "cross_aspects":[...], "house_overlay_owner":[...],
    "house_overlay_pet":[...], "sections":{...}}. owner_cover_path/pet_cover_path: fotos
    (ja estilizadas ou originais) para a capa; se None, a coluna correspondente da capa
    mostra a roda natal daquele sujeito no lugar da foto."""
    tmp_dir = tempfile.mkdtemp(prefix="report_pet_")
    owner, pet = data["owner"], data["pet"]

    wheel_owner_png = os.path.join(tmp_dir, "wheel_owner.png")
    build_wheel_png(owner["planets"], owner["houses"], owner["ascendant"], wheel_owner_png)
    wheel_pet_png = os.path.join(tmp_dir, "wheel_pet.png")
    build_wheel_png(pet["planets"], pet["houses"], pet["ascendant"], wheel_pet_png)

    data = dict(data)
    data["_wheel_png_owner"] = wheel_owner_png
    data["_wheel_png_pet"] = wheel_pet_png

    cover_image_path = os.path.join(tmp_dir, "cover.png")
    build_pet_synastry_cover_png(data, cover_image_path,
                                  owner_photo_path=owner_cover_path, pet_photo_path=pet_cover_path)

    owner_name, pet_name = owner["name"], pet["name"]
    running = f"SINASTRIA COM PET • {owner_name.upper()} & {pet_name.upper()}"

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(9.2); normal.font.color.rgb = BODY_INK

    sec0 = doc.sections[0]
    sec0.page_width = Cm(21.0); sec0.page_height = Cm(29.7)
    sec0.left_margin = sec0.right_margin = sec0.top_margin = sec0.bottom_margin = Cm(0)
    sec0.header_distance = sec0.footer_distance = Cm(0)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.add_run().add_picture(cover_image_path, width=Cm(21.0), height=Cm(29.7))

    sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec1.page_width = Cm(21.0); sec1.page_height = Cm(29.7)
    sec1.left_margin = sec1.right_margin = Cm(2.35)
    sec1.top_margin = sec1.bottom_margin = Cm(1.55)
    sec1.header_distance = Cm(0.5); sec1.footer_distance = Cm(0.6)
    sec1.header.is_linked_to_previous = False
    sec1.footer.is_linked_to_previous = False

    hp = sec1.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run(running); hr.font.name = FONT; hr.font.size = Pt(6.5); hr.font.bold = True; hr.font.color.rgb = GOLD
    para_bottom_border(hp, color=GOLD_LINE, sz=5)

    fp = sec1.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(running + " • "); fr.font.name = FONT; fr.font.size = Pt(6.3); fr.font.color.rgb = FOOTER_INK
    add_page_field(fp)

    pages = build_pet_synastry_pages(data)
    for i, pg in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        render_head(doc, pg['eyebrow'], pg['title'], pg['subtitle'])
        render_blocks(doc, pg['blocks'])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
